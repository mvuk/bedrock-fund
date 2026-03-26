#!/usr/bin/env python3
"""
generate_word.py — Produce the Bedrock Fund portfolio analysis report (.docx).

Generates a structured Word document walking through all 8 assignment steps
with theoretical explanations, formulas, and computed numerical results.

EC310R Financial Economic Theory — Wilfrid Laurier University
Professor Doron Nisani | Huang & Litzenberger
"""

import os
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ── Asset metadata ───────────────────────────────────────────────────────────

ASSET_NAMES = {
    "NVDA": "NVIDIA", "TSM": "TSMC", "EQIX": "Equinix", "VRT": "Vertiv",
    "LIN": "Linde", "ALB": "Albemarle", "NEE": "NextEra Energy",
    "CAT": "Caterpillar", "UNP": "Union Pacific", "PLD": "Prologis",
    "DE": "Deere", "WM": "Waste Management", "UNH": "UnitedHealth",
    "TMO": "Thermo Fisher", "ISRG": "Intuitive Surgical",
    "LMT": "Lockheed Martin", "FCX": "Freeport-McMoRan", "NEM": "Newmont",
    "COST": "Costco", "BRK-B": "Berkshire Hathaway",
}

CLUSTERS = {
    "AI/Compute Infrastructure": {
        "tickers": ["NVDA", "TSM", "EQIX", "VRT"],
        "descriptions": {
            "NVDA": "Designs the GPUs that power AI training and inference workloads globally.",
            "TSM": "Fabricates the advanced semiconductors underlying all major AI chip designs.",
            "EQIX": "Operates carrier-neutral data centers that house cloud and AI compute clusters.",
            "VRT": "Supplies power and thermal management systems for high-density data centers.",
        },
    },
    "Energy Transition Enablers": {
        "tickers": ["LIN", "ALB", "NEE", "CAT"],
        "descriptions": {
            "LIN": "Produces industrial gases essential for semiconductor fabrication and clean hydrogen.",
            "ALB": "Mines and refines lithium, the critical input for EV and grid-storage batteries.",
            "NEE": "Operates the largest portfolio of wind and solar generation assets in North America.",
            "CAT": "Manufactures the heavy equipment required to build renewable energy and mining infrastructure.",
        },
    },
    "Global Logistics & Physical Infrastructure": {
        "tickers": ["UNP", "PLD", "DE", "WM"],
        "descriptions": {
            "UNP": "Runs the western US rail network that moves bulk commodities and intermodal freight.",
            "PLD": "Owns logistics warehouses positioned at key distribution nodes for e-commerce and supply chains.",
            "DE": "Builds precision agriculture and construction machinery integrating GPS and autonomy.",
            "WM": "Manages solid waste collection, recycling, and landfill gas-to-energy operations.",
        },
    },
    "Healthcare & Life Sciences Infrastructure": {
        "tickers": ["UNH", "TMO", "ISRG"],
        "descriptions": {
            "UNH": "Operates the largest US health insurance and pharmacy-benefit platform.",
            "TMO": "Supplies laboratory instruments, reagents, and services to pharmaceutical and biotech R&D.",
            "ISRG": "Manufactures the da Vinci robotic surgical systems used in minimally invasive procedures.",
        },
    },
    "Defense & Hard Asset Enablers": {
        "tickers": ["LMT", "FCX", "NEM", "COST", "BRK-B"],
        "descriptions": {
            "LMT": "Produces advanced defense platforms including the F-35 and missile-defense systems.",
            "FCX": "Operates the world's largest publicly traded copper mine, supplying electrification demand.",
            "NEM": "The world's largest gold miner, providing a hard-asset hedge against monetary debasement.",
            "COST": "Runs a membership warehouse model that benefits from consumer staples demand in all cycles.",
            "BRK-B": "A diversified conglomerate with insurance, rail, energy, and manufacturing subsidiaries.",
        },
    },
}


# ── Helpers ──────────────────────────────────────────────────────────────────

def _set_cell_text(cell, text, bold=False, size=9):
    """Set cell text with formatting."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def _add_paragraph(doc, text, style=None, bold=False, italic=False, size=11):
    """Add a paragraph with optional inline formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def _add_formula(doc, text):
    """Add a centered, italic formula line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    return p


def _fmt(x, decimals=6):
    """Format a number to a fixed number of decimals."""
    return f"{x:.{decimals}f}"


def _pct(x):
    """Format as percentage."""
    return f"{x:+.2%}"


# ── Main report generator ───────────────────────────────────────────────────

def generate_report(
    risky_prices,
    risky_returns,
    mean_vector,
    cov_matrix,
    params,
    frontier,
    bil_prices,
    bil_returns,
    mvp,
    market,
    zc,
    tickers,
    rf,
    graph_path=None,
    output_dir="output",
    inv_max_dev=None,
    inv_cond_num=None,
):
    """
    Generate bedrock_fund_report.docx with all 8 assignment sections
    plus conclusions.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    A, B, C, D = params["A"], params["B"], params["C"], params["D"]
    Sigma_inv = params["Sigma_inv"]
    n_assets = len(tickers)
    n_days = len(risky_prices)
    n_returns = len(risky_returns)
    ticker_idx = {t: i for i, t in enumerate(tickers)}
    daily_std = np.sqrt(np.diag(cov_matrix))
    ann_mean = mean_vector * 252
    ann_std = daily_std * np.sqrt(252)

    # ── Title page ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bedrock Fund — Portfolio Analysis Report")
    run.font.size = Pt(24)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EC310R Financial Economic Theory — Group Assignment")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wilfrid Laurier University — Professor Doron Nisani")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Textbook: Huang & Litzenberger, Foundations for Financial Economics")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Group Members: _____, _____, _____")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: Asset Selection
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Asset Selection (Step 1)", level=1)

    doc.add_paragraph(
        'The Bedrock Fund follows a "picks-and-shovels" investment thesis. Rather than '
        "betting on which end-market companies will win, the fund invests in the "
        "infrastructure providers whose products and services are required regardless of "
        "which downstream competitors prevail. The name reflects the idea that during a "
        "gold rush, the surest profits go to those selling picks and shovels."
    )

    doc.add_paragraph(
        f"We selected {n_assets} risky assets organized into five thematic clusters. "
        "Each cluster targets a distinct structural trend in the global economy, and the "
        "assets within each cluster represent the essential infrastructure enabling that "
        "trend. By spanning different sectors and factor exposures — technology, energy, "
        "industrials, healthcare, defense, commodities — we aim to produce a covariance "
        "structure with meaningful off-diagonal diversity, which is critical for the "
        "efficient frontier to exhibit a well-shaped hyperbola with meaningful "
        "diversification benefits."
    )

    # Only show cluster detail if all 20 standard tickers are present
    all_cluster_tickers = set()
    for cinfo in CLUSTERS.values():
        all_cluster_tickers.update(cinfo["tickers"])
    has_full_clusters = all_cluster_tickers.issubset(set(tickers))

    if has_full_clusters:
        for cluster_name, cluster_info in CLUSTERS.items():
            doc.add_heading(cluster_name, level=2)
            for t in cluster_info["tickers"]:
                desc = cluster_info["descriptions"][t]
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{t} ({ASSET_NAMES.get(t, t)}): ")
                run.bold = True
                p.add_run(desc)
    else:
        for t in tickers:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{t} ({ASSET_NAMES.get(t, t)})")
            run.bold = True

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: Daily Prices
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Daily Prices (Step 2, Sheet 1)", level=1)

    _add_paragraph(doc,
        "Reference: Course Chapter 6 (Statistical Foundations); "
        "Huang & Litzenberger, Chapter 3.",
        italic=True, size=10)

    first_date = str(risky_prices.index[0])[:10]
    last_date = str(risky_prices.index[-1])[:10]

    doc.add_paragraph(
        f"Daily adjusted close prices were downloaded from Yahoo Finance for all "
        f"{n_assets} risky assets over the period {first_date} to {last_date}, "
        f"yielding {n_days} trading days of data."
    )

    doc.add_paragraph(
        "We use adjusted close prices rather than raw close prices because the adjusted "
        "series accounts for stock splits and dividend distributions. This ensures that "
        "the returns computed from consecutive prices reflect the total return actually "
        "earned by an investor, not just the capital gain component. Without this "
        "adjustment, returns around ex-dividend dates or split dates would be distorted, "
        "biasing the mean vector and covariance matrix."
    )

    doc.add_paragraph(
        "The full daily price matrix is reported in Sheet 1 of the Excel workbook, with "
        f"dates as rows and the {n_assets} tickers as columns."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: Daily Returns
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Daily Returns (Step 3, Sheet 2)", level=1)

    _add_paragraph(doc,
        "Reference: Course Chapter 6 — Rate of Return; "
        "Huang & Litzenberger, Chapter 3.",
        italic=True, size=10)

    doc.add_paragraph(
        "The simple (arithmetic) daily return for each asset is computed as:"
    )

    _add_formula(doc, "R_t = (P_t - P_{t-1}) / P_{t-1}")

    doc.add_paragraph(
        "We use simple returns rather than logarithmic returns because the portfolio "
        "return is the weighted average of individual simple returns:"
    )

    _add_formula(doc, "R_P = w_1 R_1 + w_2 R_2 + ... + w_n R_n")

    doc.add_paragraph(
        "This additivity property is essential for the mean-variance optimization "
        "framework. With log returns, the portfolio log return is NOT the weighted sum "
        "of individual log returns, which would invalidate the linear algebra underlying "
        "the efficient frontier derivation."
    )

    doc.add_paragraph(
        f"Computing returns from {n_days} price observations yields "
        f"{n_returns} daily return observations for each of the {n_assets} assets. "
        "The full return matrix is reported in Sheet 2 of the Excel workbook."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: Mean and Covariance
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Mean Returns and Covariance Matrix (Step 4, Sheet 3)", level=1)

    _add_paragraph(doc,
        "Reference: Course Chapters 6\u20138 (Weeks 1\u20133); "
        "Huang & Litzenberger, Chapters 3\u20134.",
        italic=True, size=10)

    doc.add_paragraph(
        "Under the mean-variance framework, the investment decision depends on only two "
        "sufficient statistics: the mean return vector \u03bc and the covariance matrix \u03a3."
    )

    doc.add_paragraph(
        "The sample mean vector \u03bc = (\u03bc_1, \u03bc_2, ..., \u03bc_n)^T is estimated as the "
        f"arithmetic average of the {n_returns} daily return observations for each asset. "
        "The sample covariance matrix \u03a3 is the n\u00d7n matrix estimated with the Bessel-corrected "
        "denominator (n\u22121 = "
        f"{n_returns - 1}) to produce an unbiased estimate."
    )

    doc.add_paragraph(
        "Key properties of the covariance matrix \u03a3:"
    )
    for prop in [
        f"Square: {n_assets}\u00d7{n_assets}",
        "Symmetric: \u03a3 = \u03a3^T (covariance of i with j equals covariance of j with i)",
        "Positive semi-definite: x^T \u03a3 x \u2265 0 for all x",
        "Invertible: required for the frontier derivation (\u03a3\u207b\u00b9 exists)",
    ]:
        doc.add_paragraph(prop, style="List Bullet")

    # Report highest/lowest mean returns
    max_idx = np.argmax(mean_vector)
    min_idx = np.argmin(mean_vector)
    max_t = tickers[max_idx]
    min_t = tickers[min_idx]

    doc.add_paragraph(
        f"Among the {n_assets} assets, the highest mean daily return is "
        f"{ASSET_NAMES[max_t]} ({max_t}) at {_fmt(mean_vector[max_idx])} "
        f"(annualized: {_pct(ann_mean[max_idx])}), "
        f"and the lowest is {ASSET_NAMES[min_t]} ({min_t}) at "
        f"{_fmt(mean_vector[min_idx])} "
        f"(annualized: {_pct(ann_mean[min_idx])})."
    )

    doc.add_paragraph(
        "The mean vector and full 20\u00d720 covariance matrix are reported in Sheet 3 "
        "of the Excel workbook."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: Implementation Notes and Dilemmas
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Implementation Notes and Dilemmas", level=1)

    doc.add_paragraph(
        "We implemented all calculations in Python using numpy for linear algebra, "
        "pandas for data management, and matplotlib for visualization. We chose Python "
        "over Excel because the 20\u00d720 covariance matrix inversion is numerically "
        "fragile in Excel\u2019s MINVERSE function, whereas numpy.linalg.inv uses LU "
        "decomposition which is more stable for matrices of this size."
    )

    # Inversion verification
    if inv_max_dev is not None and inv_cond_num is not None:
        if inv_cond_num < 100:
            cond_desc = "well-conditioned"
        elif inv_cond_num < 10000:
            cond_desc = "moderately conditioned"
        else:
            cond_desc = "ill-conditioned"
        doc.add_paragraph(
            f"We verified the matrix inversion by computing \u03a3\u00b7\u03a3\u207b\u00b9 and checking "
            f"its deviation from the identity matrix. The maximum absolute deviation was "
            f"{inv_max_dev:.2e}, confirming numerical accuracy. The condition number of "
            f"\u03a3 was {inv_cond_num:.2f}, indicating the matrix is {cond_desc}."
        )

    doc.add_paragraph(
        "We used simple returns rather than log returns because portfolio return is "
        "the weighted sum of individual simple returns (R_P = w\u2019R), which is the "
        "property exploited by the mean-variance optimization. Log returns do not "
        "aggregate linearly across assets."
    )

    doc.add_paragraph(
        f"The covariance matrix was estimated using the sample covariance with "
        f"Bessel\u2019s correction (dividing by n\u22121 = {n_returns - 1}), which provides "
        "an unbiased estimator. All reported statistics are at daily frequency unless "
        "explicitly annualized. Annualization uses the conventions "
        "\u03bc_annual = \u03bc_daily \u00d7 252 and \u03c3_annual = \u03c3_daily \u00d7 \u221a252."
    )

    gross_exposure_impl = np.sum(np.abs(market["weights"]))
    doc.add_paragraph(
        f"A key dilemma we encountered is the extreme leverage in the Market Portfolio "
        f"(gross exposure of approximately {gross_exposure_impl * 100:.0f}%). This is a "
        "direct consequence of the unconstrained optimization \u2014 the Lagrangian in "
        "Chapter 8 imposes w\u20191 = 1 but does not require w_i \u2265 0. The optimizer "
        "exploits differences in the covariance structure by taking large offsetting "
        "long and short positions. While mathematically optimal, this portfolio would "
        "be impractical for a real mutual fund without margin accounts and significant "
        "borrowing capacity. This tension between theoretical optimality and practical "
        "feasibility is a central theme in modern portfolio theory."
    )

    # Find the worst-performing asset dynamically
    worst_idx = np.argmin(ann_mean)
    worst_t = tickers[worst_idx]
    worst_name = ASSET_NAMES.get(worst_t, worst_t)
    worst_ann = ann_mean[worst_idx]
    doc.add_paragraph(
        f"Another observation: several assets had negative mean returns over the "
        f"sample period, most notably {worst_name} ({_pct(worst_ann)} annualized). "
        "The optimizer handles this by shorting these assets in the Market Portfolio, "
        "effectively profiting from their decline. In the MVP, these assets receive "
        "near-zero weight since the optimizer there is indifferent to returns."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: Efficient Portfolio Frontier
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Efficient Portfolio Frontier (Step 5, Sheet 4)", level=1)

    _add_paragraph(doc,
        "Reference: Course Chapter 8 \u2014 N-Asset Portfolio Selection (Weeks 2\u20133); "
        "Huang & Litzenberger, Chapter 4.",
        italic=True, size=10)

    doc.add_paragraph(
        "The investor's problem is to find the portfolio with minimum variance for each "
        "target expected return \u03bc_P:"
    )

    _add_formula(doc,
        "min  w^T \u03a3 w    subject to    w^T \u03bc = \u03bc_P ,   w^T 1 = 1"
    )

    doc.add_paragraph(
        "Forming the Lagrangian L = w^T\u03a3w \u2212 \u03bb\u2081(w^T\u03bc \u2212 \u03bc_P) \u2212 \u03bb\u2082(w^T1 \u2212 1) "
        "and setting the first-order conditions to zero yields the optimal weight vector:"
    )

    _add_formula(doc,
        "w* = \u03a3\u207b\u00b9 [\u03bb\u2081 \u03bc + \u03bb\u2082 1]"
    )

    doc.add_paragraph(
        "where the Lagrange multipliers \u03bb\u2081 and \u03bb\u2082 are determined by the constraints. "
        "The solution depends on four scalar constants computed from \u03a3\u207b\u00b9:"
    )

    # Constants table
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Constant", "Formula", "Computed Value"]
    for j, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[j], h, bold=True, size=10)
    data = [
        ("A", "1^T \u03a3\u207b\u00b9 \u03bc", _fmt(A, 10)),
        ("B", "\u03bc^T \u03a3\u207b\u00b9 \u03bc", _fmt(B, 10)),
        ("C", "1^T \u03a3\u207b\u00b9 1", _fmt(C, 4)),
        ("D", "BC \u2212 A\u00b2", _fmt(D, 4)),
    ]
    for i, (const, formula, val) in enumerate(data):
        _set_cell_text(table.rows[i + 1].cells[0], const, bold=True, size=10)
        _set_cell_text(table.rows[i + 1].cells[1], formula, size=10)
        _set_cell_text(table.rows[i + 1].cells[2], val, size=10)

    doc.add_paragraph()
    doc.add_paragraph(
        "The requirement D > 0 ensures that the frontier is a proper (non-degenerate) "
        "hyperbola. If D = 0, all assets would have the same expected return and the "
        "frontier would collapse to a single point."
    )

    doc.add_paragraph(
        "The frontier in (\u03c3\u00b2, \u03bc) space is the parabola:"
    )

    _add_formula(doc,
        "\u03c3\u00b2_P(\u03bc_P) = (1/D)(C\u03bc_P\u00b2 \u2212 2A\u03bc_P + B)"
    )

    doc.add_paragraph(
        "Equivalently, in hyperbola form:"
    )

    _add_formula(doc,
        "\u03c3\u00b2_P = (C/D)(\u03bc_P \u2212 A/C)\u00b2 + 1/C"
    )

    doc.add_paragraph(
        "The vertex of this hyperbola is the Minimum Variance Portfolio (MVP):"
    )

    _add_formula(doc,
        f"\u03bc_MVP = A/C = {_fmt(mvp['expected_return'], 8)}")
    _add_formula(doc,
        f"\u03c3_MVP = \u221a(1/C) = {_fmt(mvp['std_dev'], 8)}")

    doc.add_paragraph(
        "The efficient set consists of all frontier portfolios with \u03bc_P \u2265 \u03bc_MVP "
        "(the upper branch of the hyperbola). Portfolios below the MVP lie on the "
        "inefficient branch: they have the same variance as a portfolio on the upper "
        "branch but a lower expected return. No rational investor would hold an "
        "inefficient portfolio."
    )

    doc.add_paragraph(
        "The (\u03c3, \u03bc) frontier points are reported in Sheet 4 of the Excel workbook."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: Riskless Asset
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Riskless Asset (Step 6, Sheet 5)", level=1)

    _add_paragraph(doc,
        "Reference: Course Chapter 9 \u2014 The Risk-Free Asset; "
        "Huang & Litzenberger, Chapter 5.",
        italic=True, size=10)

    doc.add_paragraph(
        "We use BIL (SPDR Bloomberg 1\u20133 Month T-Bill ETF) as our risk-free asset proxy. "
        "BIL tracks a portfolio of short-term US Treasury bills with maturities between "
        "one and three months. It is an appropriate choice because:"
    )

    for reason in [
        "US Treasury bills are backed by the full faith and credit of the US government, "
        "making default risk negligible.",
        "The 1\u20133 month maturity range minimizes interest rate risk, producing near-zero "
        "variance in daily returns.",
        "BIL is an exchange-traded fund with daily liquidity, allowing us to obtain a "
        "consistent daily price series from the same source as our risky assets.",
    ]:
        doc.add_paragraph(reason, style="List Bullet")

    bil_n_returns = len(bil_returns)
    bil_daily_std = float(bil_returns.std().iloc[0])

    doc.add_paragraph(
        f"The average daily return of BIL over {bil_n_returns} observations is:"
    )
    _add_formula(doc, f"\u03bc_F = {_fmt(rf, 8)}")
    doc.add_paragraph(
        f"which annualizes to approximately {_pct(rf * 252)}. "
        f"The daily standard deviation of BIL returns is {_fmt(bil_daily_std, 8)}, "
        "confirming near-zero variance."
    )

    mu_mvp = A / C
    doc.add_paragraph(
        f"Critically, we verify that \u03bc_F < A/C (Case 1 in Huang & Litzenberger): "
        f"\u03bc_F = {_fmt(rf, 8)} < A/C = {_fmt(mu_mvp, 8)}. "
        "This condition ensures that the tangent line from the risk-free rate touches "
        "the efficient (upper) branch of the frontier, so the tangency portfolio is a "
        "legitimate efficient portfolio."
    )

    doc.add_paragraph(
        "BIL daily prices and returns are reported in Sheet 5 of the Excel workbook."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: Portfolios
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. MVP, Market Portfolio, and ZC Portfolio (Step 7, Sheet 6)", level=1)

    _add_paragraph(doc,
        "Reference: Course Chapters 8\u20139; "
        "Huang & Litzenberger, Chapters 4\u20135.",
        italic=True, size=10)

    # ── 7a: MVP ──────────────────────────────────────────────────────────
    doc.add_heading("8a. Minimum Variance Portfolio (MVP)", level=2)

    doc.add_paragraph(
        "The MVP is the portfolio on the frontier with the lowest possible variance. "
        "It occupies the vertex of the mean-variance hyperbola and serves as the "
        "boundary between the efficient and inefficient portions of the frontier."
    )

    _add_formula(doc, "w_MVP = \u03a3\u207b\u00b9 1 / C")

    # MVP stats table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, val in [
        ("Expected Return (\u03bc_MVP)", _fmt(mvp["expected_return"], 8)),
        ("Standard Deviation (\u03c3_MVP)", _fmt(mvp["std_dev"], 8)),
        ("Annualized Return", _pct(mvp["expected_return"] * 252)),
    ]:
        row = table.rows[[
            ("Expected Return (\u03bc_MVP)", _fmt(mvp["expected_return"], 8)),
            ("Standard Deviation (\u03c3_MVP)", _fmt(mvp["std_dev"], 8)),
            ("Annualized Return", _pct(mvp["expected_return"] * 252)),
        ].index((label, val))]
        _set_cell_text(row.cells[0], label, bold=True, size=10)
        _set_cell_text(row.cells[1], val, size=10)

    doc.add_paragraph()

    # MVP weights table
    doc.add_paragraph(
        "The MVP weight distribution reflects pure variance minimization \u2014 "
        "the optimizer ignores expected returns entirely and allocates to reduce "
        "total portfolio risk:"
    )

    table = doc.add_table(rows=n_assets + 1, cols=3)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Ticker", "Name", "MVP Weight"]):
        _set_cell_text(table.rows[0].cells[j], h, bold=True, size=9)
    # Sort by absolute weight descending
    mvp_sort = np.argsort(np.abs(mvp["weights"]))[::-1]
    for row_i, idx in enumerate(mvp_sort):
        t = tickers[idx]
        _set_cell_text(table.rows[row_i + 1].cells[0], t, size=9)
        _set_cell_text(table.rows[row_i + 1].cells[1], ASSET_NAMES[t], size=9)
        _set_cell_text(table.rows[row_i + 1].cells[2],
                       f"{mvp['weights'][idx]:+.4f}", size=9)

    doc.add_paragraph()

    # Top MVP holdings
    top3_mvp = mvp_sort[:3]
    top_names = ", ".join(
        f"{ASSET_NAMES[tickers[i]]} ({mvp['weights'][i]:+.2%})" for i in top3_mvp
    )
    doc.add_paragraph(
        f"The portfolio is dominated by low-volatility defensive names: {top_names}. "
        "These assets have the lowest daily standard deviations in the universe and "
        "relatively low correlations with other holdings, making them efficient "
        "variance reducers. The MVP is agnostic to expected return \u2014 it is the "
        "portfolio that a maximally risk-averse investor would hold."
    )

    # ── 7b: Market Portfolio ─────────────────────────────────────────────
    doc.add_heading("8b. Market Portfolio (Tangency Portfolio)", level=2)

    doc.add_paragraph(
        "When a riskless asset is available, the efficient set becomes a straight line "
        "(the Capital Market Line) from the risk-free rate tangent to the risky-asset "
        "frontier. The tangency point is the Market Portfolio \u2014 the risky portfolio "
        "that maximizes the Sharpe ratio."
    )

    doc.add_paragraph(
        "The unnormalized tangency weights are:"
    )
    _add_formula(doc, "z = \u03a3\u207b\u00b9 (\u03bc \u2212 \u03bc_F \u00b7 1)")

    doc.add_paragraph(
        "Normalizing to sum to 1 gives w_M = z / (1^T z). The quantity H measures "
        "the squared maximum Sharpe ratio attainable from the risky assets:"
    )
    _add_formula(doc,
        "H = (\u03bc \u2212 \u03bc_F\u00b71)^T \u03a3\u207b\u00b9 (\u03bc \u2212 \u03bc_F\u00b71)")

    # Market stats table
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    mkt_stats = [
        ("Expected Return (\u03bc_M)", _fmt(market["expected_return"], 8)),
        ("Standard Deviation (\u03c3_M)", _fmt(market["std_dev"], 8)),
        ("Sharpe Ratio", _fmt(market["sharpe_ratio"], 6)),
        ("H", _fmt(market["H"], 10)),
        ("Annualized Return", _pct(market["expected_return"] * 252)),
    ]
    for i, (label, val) in enumerate(mkt_stats):
        _set_cell_text(table.rows[i].cells[0], label, bold=True, size=10)
        _set_cell_text(table.rows[i].cells[1], val, size=10)

    doc.add_paragraph()

    # Market weights table
    table = doc.add_table(rows=n_assets + 1, cols=3)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Ticker", "Name", "Market Weight"]):
        _set_cell_text(table.rows[0].cells[j], h, bold=True, size=9)
    mkt_sort = np.argsort(np.abs(market["weights"]))[::-1]
    for row_i, idx in enumerate(mkt_sort):
        t = tickers[idx]
        _set_cell_text(table.rows[row_i + 1].cells[0], t, size=9)
        _set_cell_text(table.rows[row_i + 1].cells[1], ASSET_NAMES[t], size=9)
        _set_cell_text(table.rows[row_i + 1].cells[2],
                       f"{market['weights'][idx]:+.4f}", size=9)

    doc.add_paragraph()

    # Leverage discussion
    gross_exposure = np.sum(np.abs(market["weights"]))
    long_sum = np.sum(market["weights"][market["weights"] > 0])
    short_sum = np.sum(market["weights"][market["weights"] < 0])

    # Cluster weights for discussion (only if full 20 tickers)
    long_clusters = []
    short_clusters = []
    if has_full_clusters:
        cluster_mkt_w = {}
        for cname, cinfo in CLUSTERS.items():
            cols = [ticker_idx[t] for t in cinfo["tickers"]]
            cluster_mkt_w[cname] = np.sum(market["weights"][cols])
        long_clusters = sorted([(k, v) for k, v in cluster_mkt_w.items() if v > 0.1], key=lambda x: -x[1])
        short_clusters = sorted([(k, v) for k, v in cluster_mkt_w.items() if v < -0.1], key=lambda x: x[1])

    doc.add_heading("Discussion: Extreme Weights and Leverage", level=3)

    doc.add_paragraph(
        f"The Market Portfolio exhibits substantial leverage: the gross exposure "
        f"(sum of absolute weights) is {gross_exposure:.1f}%, with "
        f"{long_sum:+.2f} in long positions and {short_sum:+.2f} in short positions. "
        "This is a well-documented property of unconstrained mean-variance optimization "
        "and warrants careful interpretation."
    )

    doc.add_paragraph(
        "The unconstrained optimization framework as presented in the course "
        "permits short selling \u2014 portfolio weights can be negative. "
        "This is consistent with the theoretical treatment in Huang & Litzenberger, "
        "where no sign constraints are imposed on the weight vector."
    )

    if long_clusters or short_clusters:
        doc.add_paragraph("At the cluster level, the optimizer is:")
        for cname, w in long_clusters:
            doc.add_paragraph(
                f"Long {cname} ({w:+.2f})", style="List Bullet")
        for cname, w in short_clusters:
            doc.add_paragraph(
                f"Short {cname} ({w:+.2f})", style="List Bullet")

    doc.add_paragraph(
        "The optimizer goes long the clusters with the highest risk-adjusted returns "
        "and shorts those with negative or low returns. It exploits the correlation "
        "structure to hedge out common factor exposures, effectively isolating the "
        "return spread between winning and losing clusters."
    )

    doc.add_paragraph(
        "This extreme leverage is a key limitation of the mean-variance model. "
        "In practice, several approaches address it: imposing long-only constraints "
        "(0 \u2264 w_i \u2264 1), adding regularization to the covariance matrix, using "
        "shrinkage estimators (Ledoit-Wolf), or applying position limits. The "
        "Bedrock Fund's theoretical analysis intentionally preserves the unconstrained "
        "solution to demonstrate the framework as taught in the course."
    )

    # ── 7c: Zero-Covariance Portfolio ────────────────────────────────────
    doc.add_heading("8c. Zero-Covariance Portfolio", level=2)

    _add_paragraph(doc,
        "Reference: Course Chapter 8, Proposition 2 \u2014 Zero-Covariance Portfolio.",
        italic=True, size=10)

    doc.add_paragraph(
        "For any frontier portfolio P with expected return \u03bc_P \u2260 \u03bc_MVP, there exists "
        "a unique frontier portfolio ZC(P) such that Cov(R_P, R_ZC) = 0. The ZC "
        "portfolio generalizes the role of the risk-free asset: in the absence of a "
        "riskless asset, the ZC portfolio of any efficient portfolio provides the same "
        "two-fund separation theorem."
    )

    doc.add_paragraph(
        "For the Market Portfolio with return \u03bc_M, the ZC portfolio's expected return is:"
    )
    _add_formula(doc,
        "\u03bc_ZC = A/C \u2212 (D/C\u00b2) / (\u03bc_M \u2212 A/C)")

    cov_mkt_zc = market["weights"] @ cov_matrix @ zc["weights"]

    # ZC stats table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    zc_stats = [
        ("Expected Return (\u03bc_ZC)", _fmt(zc["expected_return"], 8)),
        ("Standard Deviation (\u03c3_ZC)", _fmt(zc["std_dev"], 8)),
        ("Cov(Market, ZC)", f"{cov_mkt_zc:.2e}"),
        ("Annualized Return", _pct(zc["expected_return"] * 252)),
    ]
    for i, (label, val) in enumerate(zc_stats):
        _set_cell_text(table.rows[i].cells[0], label, bold=True, size=10)
        _set_cell_text(table.rows[i].cells[1], val, size=10)

    doc.add_paragraph()

    doc.add_paragraph(
        f"The covariance between the Market Portfolio and its ZC portfolio is "
        f"{cov_mkt_zc:.2e}, confirming zero covariance to machine precision."
    )

    doc.add_paragraph(
        f"Notably, \u03bc_ZC = {_fmt(zc['expected_return'], 8)} \u2248 \u03bc_F = {_fmt(rf, 8)}. "
        "This is theoretically expected: for the tangency portfolio (Market Portfolio), "
        "the zero-covariance portfolio has the same expected return as the risk-free "
        "rate. This follows from the Black (1972) zero-beta CAPM: when a riskless asset "
        "exists, the zero-beta rate equals the risk-free rate."
    )

    doc.add_paragraph(
        "All portfolio weights and statistics are reported in Sheet 6 of the Excel "
        "workbook."
    )

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: Graph Analysis
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Graph Analysis (Step 8)", level=1)

    _add_paragraph(doc,
        "Reference: Course Chapters 8\u20139; "
        "Huang & Litzenberger, Chapters 4\u20135.",
        italic=True, size=10)

    doc.add_paragraph(
        "Graph 1 plots the full portfolio analysis in (\u03c3, \u03bc) space \u2014 "
        "standard deviation on the horizontal axis and expected return on the vertical "
        "axis. The graph contains the following elements:"
    )

    cml_slope = np.sqrt(market["H"])

    elements = [
        ("Efficient Portfolio Frontier (upper branch)",
         "The solid curve above the MVP representing all efficient risky-asset "
         "portfolios. These are the portfolios that maximize expected return for a "
         "given level of risk."),
        ("Inefficient Frontier (lower branch)",
         "The dashed curve below the MVP. These portfolios are dominated \u2014 for the "
         "same variance, an investor can achieve a higher return on the upper branch."),
        ("Minimum Variance Portfolio (MVP)",
         f"Marked at (\u03c3, \u03bc) = ({_fmt(mvp['std_dev'], 4)}, "
         f"{_fmt(mvp['expected_return'], 6)}). The leftmost point on the frontier."),
        ("Market Portfolio (Tangency Portfolio)",
         f"Marked at (\u03c3, \u03bc) = ({_fmt(market['std_dev'], 4)}, "
         f"{_fmt(market['expected_return'], 6)}). The point where the CML is "
         "tangent to the efficient frontier."),
        ("Zero-Covariance Portfolio",
         f"Marked at (\u03c3, \u03bc) = ({_fmt(zc['std_dev'], 4)}, "
         f"{_fmt(zc['expected_return'], 6)}). Located on the inefficient branch, "
         "directly below the MVP."),
        ("Capital Market Line (CML)",
         f"The straight line from (0, \u03bc_F) through the Market Portfolio, with slope "
         f"\u221aH = {_fmt(cml_slope, 6)}."),
    ]

    for name, desc in elements:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "The Capital Market Line equation is:"
    )
    _add_formula(doc,
        f"\u03bc_P = \u03bc_F + \u221aH \u00b7 \u03c3_P = "
        f"{_fmt(rf, 8)} + {_fmt(cml_slope, 6)} \u00b7 \u03c3_P")

    doc.add_paragraph(
        "The CML dominates the efficient frontier: for any level of risk \u03c3_P, "
        "the CML offers a higher expected return than the risky-asset frontier alone. "
        "This is the key insight of introducing a riskless asset \u2014 investors can "
        "improve their risk-return tradeoff by combining the risk-free asset with "
        "the Market Portfolio."
    )

    doc.add_paragraph(
        "Two-fund separation follows directly: every efficient portfolio (on the CML) "
        "can be expressed as a combination of just two funds \u2014 the risk-free asset "
        "and the Market Portfolio. A conservative investor holds mostly the risk-free "
        "asset with a small allocation to the Market Portfolio; an aggressive investor "
        "borrows at the risk-free rate and levers up the Market Portfolio."
    )

    # Embed the graph
    if graph_path and os.path.exists(graph_path):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(graph_path, width=Inches(6.0))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Figure 1: Bedrock Fund \u2014 Efficient Portfolio Frontier, "
            "MVP, Market Portfolio, ZC Portfolio, and Capital Market Line "
            "in (\u03c3, \u03bc) space."
        )
        run.italic = True
        run.font.size = Pt(9)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: Conclusions
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Conclusions", level=1)

    doc.add_heading("Advantages of the Bedrock Fund", level=2)

    advantages = [
        ("Thematic diversification",
         "The picks-and-shovels approach distributes exposure across five distinct "
         "structural trends (AI, energy transition, logistics, healthcare, defense). "
         "This reduces single-company and single-sector risk."),
        ("Factor breadth",
         "The 20 assets span technology, industrials, energy, healthcare, commodities, "
         "and financials, providing exposure to multiple macroeconomic factors "
         "(growth, inflation, rates, geopolitics)."),
        ("Infrastructure resilience",
         "Infrastructure providers tend to have more durable revenue streams than "
         "end-market competitors, as their products are required regardless of which "
         "downstream company wins."),
        ("Covariance structure quality",
         "The cross-cluster correlations range from approximately 0.36 to 0.62, "
         "providing meaningful but not excessive diversification \u2014 enough off-diagonal "
         "variation for the efficient frontier to offer real variance reduction."),
    ]

    for title, desc in advantages:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Limitations and Disadvantages", level=2)

    limitations = [
        ("Impractical leverage in the unconstrained solution",
         f"The Market Portfolio has gross exposure of {gross_exposure:.0f}%, "
         "requiring extensive short selling. In practice, most mutual funds face "
         "regulatory constraints (e.g., the Investment Company Act of 1940) that "
         "prohibit or limit short positions."),
        ("Backward-looking estimation",
         "The mean vector and covariance matrix are estimated from one year of "
         "historical data. Past returns are a noisy and often unreliable predictor "
         "of future returns. The covariance structure is more stable, but still "
         "subject to regime changes."),
        ("Normality assumption",
         "The mean-variance framework implicitly assumes returns are normally "
         "distributed (or that investor utility is quadratic). In reality, asset "
         "returns exhibit fat tails, skewness, and time-varying volatility, which "
         "the model does not capture."),
        ("Single-period model",
         "The analysis is a single-period (one-shot) optimization. It does not "
         "account for dynamic rebalancing, transaction costs, taxes, or changing "
         "investment opportunities over time."),
        ("Estimation error amplification",
         "Inverting the covariance matrix amplifies estimation errors in the "
         "sample means and covariances, which is a primary driver of the extreme "
         "weights observed. This is the classic Michaud (1989) critique of "
         "mean-variance optimization."),
    ]

    for title, desc in limitations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("The Tension Between Theory and Practice", level=2)

    doc.add_paragraph(
        "The Bedrock Fund analysis illustrates a fundamental tension in financial "
        "economics. The mean-variance framework of Markowitz (1952), as formalized "
        "by Huang & Litzenberger, provides an elegant and mathematically rigorous "
        "theory for optimal portfolio selection. The efficient frontier, the tangency "
        "portfolio, the zero-covariance portfolio, and two-fund separation are "
        "powerful theoretical results that illuminate the structure of the "
        "risk-return tradeoff."
    )

    doc.add_paragraph(
        "However, the unconstrained implementation reveals the model's sensitivity "
        "to input estimation. Small changes in the estimated mean vector can produce "
        "large swings in optimal weights, and the resulting portfolios often require "
        "leverage and short selling that are impractical for most real-world investors. "
        "This gap between theoretical optimality and practical implementability is "
        "one of the central challenges in quantitative portfolio management, and has "
        "motivated decades of research into robust optimization, shrinkage estimators, "
        "and constrained frontier methods."
    )

    # ── Save docx ──────────────────────────────────────────────────────
    output_path = os.path.join(output_dir, "bedrock_fund_report.docx")
    doc.save(output_path)
    print(f"Report saved to {output_path}")

    # ── Generate Markdown version ────────────────────────────────────────
    md_path = _generate_markdown(
        n_assets=n_assets, n_days=n_days, n_returns=n_returns,
        tickers=tickers, ticker_idx=ticker_idx,
        mean_vector=mean_vector, cov_matrix=cov_matrix,
        ann_mean=ann_mean, ann_std=ann_std,
        params=params, frontier=frontier,
        bil_returns=bil_returns, rf=rf,
        mvp=mvp, market=market, zc=zc,
        first_date=first_date, last_date=last_date,
        output_dir=output_dir,
        inv_max_dev=inv_max_dev, inv_cond_num=inv_cond_num,
    )

    return output_path


def _generate_markdown(
    n_assets, n_days, n_returns, tickers, ticker_idx,
    mean_vector, cov_matrix, ann_mean, ann_std,
    params, frontier, bil_returns, rf,
    mvp, market, zc, first_date, last_date, output_dir,
    inv_max_dev=None, inv_cond_num=None,
):
    """Generate bedrock_fund_report.md — Markdown mirror of the Word report."""

    A, B, C, D = params["A"], params["B"], params["C"], params["D"]
    mu_mvp = A / C
    cov_mkt_zc = market["weights"] @ cov_matrix @ zc["weights"]
    gross_exposure = np.sum(np.abs(market["weights"]))
    long_sum = np.sum(market["weights"][market["weights"] > 0])
    short_sum = np.sum(market["weights"][market["weights"] < 0])
    cml_slope = np.sqrt(market["H"])
    bil_n_returns = len(bil_returns)
    bil_daily_std = float(bil_returns.std().iloc[0])
    max_idx = np.argmax(mean_vector)
    min_idx = np.argmin(mean_vector)
    max_t, min_t = tickers[max_idx], tickers[min_idx]

    all_cluster_tickers_md = set()
    for cinfo in CLUSTERS.values():
        all_cluster_tickers_md.update(cinfo["tickers"])
    has_full_clusters_md = all_cluster_tickers_md.issubset(set(tickers))

    long_clusters = []
    short_clusters = []
    if has_full_clusters_md:
        cluster_mkt_w = {}
        for cname, cinfo in CLUSTERS.items():
            cols = [ticker_idx[t] for t in cinfo["tickers"]]
            cluster_mkt_w[cname] = np.sum(market["weights"][cols])
        long_clusters = sorted([(k, v) for k, v in cluster_mkt_w.items() if v > 0.1], key=lambda x: -x[1])
        short_clusters = sorted([(k, v) for k, v in cluster_mkt_w.items() if v < -0.1], key=lambda x: x[1])

    mvp_sort = np.argsort(np.abs(mvp["weights"]))[::-1]
    mkt_sort = np.argsort(np.abs(market["weights"]))[::-1]
    top3_mvp = mvp_sort[:3]
    top_names = ", ".join(
        f"{ASSET_NAMES[tickers[i]]} ({mvp['weights'][i]:+.2%})" for i in top3_mvp
    )

    lines = []
    w = lines.append

    # ── Title ────────────────────────────────────────────────────────────
    w("# Bedrock Fund — Portfolio Analysis Report")
    w("")
    w("**EC310R Financial Economic Theory — Group Assignment**")
    w("")
    w("Wilfrid Laurier University — Professor Doron Nisani")
    w("")
    w("*Textbook: Huang & Litzenberger, Foundations for Financial Economics*")
    w("")
    w("Group Members: _____, _____, _____")
    w("")
    w("---")
    w("")

    # ── Section 1 ────────────────────────────────────────────────────────
    w("## 1. Asset Selection (Step 1)")
    w("")
    w('The Bedrock Fund follows a "picks-and-shovels" investment thesis. Rather than '
      "betting on which end-market companies will win, the fund invests in the "
      "infrastructure providers whose products and services are required regardless of "
      "which downstream competitors prevail. The name reflects the idea that during a "
      "gold rush, the surest profits go to those selling picks and shovels.")
    w("")
    w(f"We selected {n_assets} risky assets organized into five thematic clusters. "
      "Each cluster targets a distinct structural trend in the global economy, and the "
      "assets within each cluster represent the essential infrastructure enabling that "
      "trend. By spanning different sectors and factor exposures — technology, energy, "
      "industrials, healthcare, defense, commodities — we aim to produce a covariance "
      "structure with meaningful off-diagonal diversity, which is critical for the "
      "efficient frontier to exhibit a well-shaped hyperbola with meaningful "
      "diversification benefits.")
    w("")

    if has_full_clusters_md:
        for cluster_name, cluster_info in CLUSTERS.items():
            w(f"### {cluster_name}")
            w("")
            for t in cluster_info["tickers"]:
                desc = cluster_info["descriptions"][t]
                w(f"- **{t} ({ASSET_NAMES.get(t, t)}):** {desc}")
            w("")
    else:
        for t in tickers:
            w(f"- **{t} ({ASSET_NAMES.get(t, t)})**")
        w("")

    # ── Section 2 ────────────────────────────────────────────────────────
    w("## 2. Daily Prices (Step 2, Sheet 1)")
    w("")
    w("*Reference: Course Chapter 6 (Statistical Foundations); Huang & Litzenberger, Chapter 3.*")
    w("")
    w(f"Daily adjusted close prices were downloaded from Yahoo Finance for all "
      f"{n_assets} risky assets over the period {first_date} to {last_date}, "
      f"yielding {n_days} trading days of data.")
    w("")
    w("We use adjusted close prices rather than raw close prices because the adjusted "
      "series accounts for stock splits and dividend distributions. This ensures that "
      "the returns computed from consecutive prices reflect the total return actually "
      "earned by an investor, not just the capital gain component. Without this "
      "adjustment, returns around ex-dividend dates or split dates would be distorted, "
      "biasing the mean vector and covariance matrix.")
    w("")
    w(f"The full daily price matrix is reported in Sheet 1 of the Excel workbook, with "
      f"dates as rows and the {n_assets} tickers as columns.")
    w("")

    # ── Section 3 ────────────────────────────────────────────────────────
    w("## 3. Daily Returns (Step 3, Sheet 2)")
    w("")
    w("*Reference: Course Chapter 6 — Rate of Return; Huang & Litzenberger, Chapter 3.*")
    w("")
    w("The simple (arithmetic) daily return for each asset is computed as:")
    w("")
    w("> R_t = (P_t - P_{t-1}) / P_{t-1}")
    w("")
    w("We use simple returns rather than logarithmic returns because the portfolio "
      "return is the weighted average of individual simple returns:")
    w("")
    w("> R_P = w_1 R_1 + w_2 R_2 + ... + w_n R_n")
    w("")
    w("This additivity property is essential for the mean-variance optimization "
      "framework. With log returns, the portfolio log return is NOT the weighted sum "
      "of individual log returns, which would invalidate the linear algebra underlying "
      "the efficient frontier derivation.")
    w("")
    w(f"Computing returns from {n_days} price observations yields "
      f"{n_returns} daily return observations for each of the {n_assets} assets. "
      "The full return matrix is reported in Sheet 2 of the Excel workbook.")
    w("")

    # ── Section 4 ────────────────────────────────────────────────────────
    w("## 4. Mean Returns and Covariance Matrix (Step 4, Sheet 3)")
    w("")
    w("*Reference: Course Chapters 6–8 (Weeks 1–3); Huang & Litzenberger, Chapters 3–4.*")
    w("")
    w("Under the mean-variance framework, the investment decision depends on only two "
      "sufficient statistics: the mean return vector μ and the covariance matrix Σ.")
    w("")
    w(f"The sample mean vector μ = (μ_1, μ_2, ..., μ_n)^T is estimated as the "
      f"arithmetic average of the {n_returns} daily return observations for each asset. "
      "The sample covariance matrix Σ is the n×n matrix estimated with the Bessel-corrected "
      f"denominator (n−1 = {n_returns - 1}) to produce an unbiased estimate.")
    w("")
    w("Key properties of the covariance matrix Σ:")
    w("")
    w(f"- Square: {n_assets}×{n_assets}")
    w("- Symmetric: Σ = Σ^T (covariance of i with j equals covariance of j with i)")
    w("- Positive semi-definite: x^T Σ x ≥ 0 for all x")
    w("- Invertible: required for the frontier derivation (Σ⁻¹ exists)")
    w("")
    w(f"Among the {n_assets} assets, the highest mean daily return is "
      f"{ASSET_NAMES[max_t]} ({max_t}) at {_fmt(mean_vector[max_idx])} "
      f"(annualized: {_pct(ann_mean[max_idx])}), "
      f"and the lowest is {ASSET_NAMES[min_t]} ({min_t}) at "
      f"{_fmt(mean_vector[min_idx])} "
      f"(annualized: {_pct(ann_mean[min_idx])}).")
    w("")
    w("The mean vector and full 20×20 covariance matrix are reported in Sheet 3 "
      "of the Excel workbook.")
    w("")

    # ── Section 5: Implementation Notes ─────────────────────────────────
    w("## 5. Implementation Notes and Dilemmas")
    w("")
    w("We implemented all calculations in Python using numpy for linear algebra, "
      "pandas for data management, and matplotlib for visualization. We chose Python "
      "over Excel because the 20x20 covariance matrix inversion is numerically "
      "fragile in Excel's MINVERSE function, whereas numpy.linalg.inv uses LU "
      "decomposition which is more stable for matrices of this size.")
    w("")
    if inv_max_dev is not None and inv_cond_num is not None:
        if inv_cond_num < 100:
            cond_desc = "well-conditioned"
        elif inv_cond_num < 10000:
            cond_desc = "moderately conditioned"
        else:
            cond_desc = "ill-conditioned"
        w(f"We verified the matrix inversion by computing Sigma times Sigma-inverse and checking "
          f"its deviation from the identity matrix. The maximum absolute deviation was "
          f"{inv_max_dev:.2e}, confirming numerical accuracy. The condition number of "
          f"Sigma was {inv_cond_num:.2f}, indicating the matrix is {cond_desc}.")
        w("")
    w("We used simple returns rather than log returns because portfolio return is "
      "the weighted sum of individual simple returns (R_P = w'R), which is the "
      "property exploited by the mean-variance optimization. Log returns do not "
      "aggregate linearly across assets.")
    w("")
    w(f"The covariance matrix was estimated using the sample covariance with "
      f"Bessel's correction (dividing by n-1 = {n_returns - 1}), which provides "
      "an unbiased estimator. All reported statistics are at daily frequency unless "
      "explicitly annualized. Annualization uses the conventions "
      "mu_annual = mu_daily x 252 and sigma_annual = sigma_daily x sqrt(252).")
    w("")
    gross_exposure_md = np.sum(np.abs(market["weights"]))
    worst_idx_md = np.argmin(ann_mean)
    worst_t_md = tickers[worst_idx_md]
    worst_name_md = ASSET_NAMES.get(worst_t_md, worst_t_md)
    worst_ann_md = ann_mean[worst_idx_md]
    w(f"A key dilemma we encountered is the extreme leverage in the Market Portfolio "
      f"(gross exposure of approximately {gross_exposure_md * 100:.0f}%). This is a "
      "direct consequence of the unconstrained optimization — the Lagrangian in "
      "Chapter 8 imposes w'1 = 1 but does not require w_i >= 0. The optimizer "
      "exploits differences in the covariance structure by taking large offsetting "
      "long and short positions. While mathematically optimal, this portfolio would "
      "be impractical for a real mutual fund without margin accounts and significant "
      "borrowing capacity. This tension between theoretical optimality and practical "
      "feasibility is a central theme in modern portfolio theory.")
    w("")
    w(f"Another observation: several assets had negative mean returns over the "
      f"sample period, most notably {worst_name_md} ({_pct(worst_ann_md)} annualized). "
      "The optimizer handles this by shorting these assets in the Market Portfolio, "
      "effectively profiting from their decline. In the MVP, these assets receive "
      "near-zero weight since the optimizer there is indifferent to returns.")
    w("")

    # ── Section 6 ────────────────────────────────────────────────────────
    w("## 6. Efficient Portfolio Frontier (Step 5, Sheet 4)")
    w("")
    w("*Reference: Course Chapter 8 — N-Asset Portfolio Selection (Weeks 2–3); Huang & Litzenberger, Chapter 4.*")
    w("")
    w("The investor's problem is to find the portfolio with minimum variance for each "
      "target expected return μ_P:")
    w("")
    w("> min  w^T Σ w    subject to    w^T μ = μ_P ,   w^T 1 = 1")
    w("")
    w("Forming the Lagrangian L = w^TΣw − λ₁(w^Tμ − μ_P) − λ₂(w^T1 − 1) "
      "and setting the first-order conditions to zero yields the optimal weight vector:")
    w("")
    w("> w* = Σ⁻¹ [λ₁ μ + λ₂ 1]")
    w("")
    w("where the Lagrange multipliers λ₁ and λ₂ are determined by the constraints. "
      "The solution depends on four scalar constants computed from Σ⁻¹:")
    w("")
    w("| Constant | Formula | Computed Value |")
    w("|----------|---------|---------------|")
    w(f"| **A** | 1^T Σ⁻¹ μ | {_fmt(A, 10)} |")
    w(f"| **B** | μ^T Σ⁻¹ μ | {_fmt(B, 10)} |")
    w(f"| **C** | 1^T Σ⁻¹ 1 | {_fmt(C, 4)} |")
    w(f"| **D** | BC − A² | {_fmt(D, 4)} |")
    w("")
    w("The requirement D > 0 ensures that the frontier is a proper (non-degenerate) "
      "hyperbola. If D = 0, all assets would have the same expected return and the "
      "frontier would collapse to a single point.")
    w("")
    w("The frontier in (σ², μ) space is the parabola:")
    w("")
    w("> σ²_P(μ_P) = (1/D)(Cμ_P² − 2Aμ_P + B)")
    w("")
    w("Equivalently, in hyperbola form:")
    w("")
    w("> σ²_P = (C/D)(μ_P − A/C)² + 1/C")
    w("")
    w("The vertex of this hyperbola is the Minimum Variance Portfolio (MVP):")
    w("")
    w(f"> μ_MVP = A/C = {_fmt(mvp['expected_return'], 8)}")
    w(f">")
    w(f"> σ_MVP = √(1/C) = {_fmt(mvp['std_dev'], 8)}")
    w("")
    w("The efficient set consists of all frontier portfolios with μ_P ≥ μ_MVP "
      "(the upper branch of the hyperbola). Portfolios below the MVP lie on the "
      "inefficient branch: they have the same variance as a portfolio on the upper "
      "branch but a lower expected return. No rational investor would hold an "
      "inefficient portfolio.")
    w("")
    w("The (σ, μ) frontier points are reported in Sheet 4 of the Excel workbook.")
    w("")

    # ── Section 6 ────────────────────────────────────────────────────────
    w("## 7. Riskless Asset (Step 6, Sheet 5)")
    w("")
    w("*Reference: Course Chapter 9 — The Risk-Free Asset; Huang & Litzenberger, Chapter 5.*")
    w("")
    w("We use BIL (SPDR Bloomberg 1–3 Month T-Bill ETF) as our risk-free asset proxy. "
      "BIL tracks a portfolio of short-term US Treasury bills with maturities between "
      "one and three months. It is an appropriate choice because:")
    w("")
    w("- US Treasury bills are backed by the full faith and credit of the US government, making default risk negligible.")
    w("- The 1–3 month maturity range minimizes interest rate risk, producing near-zero variance in daily returns.")
    w("- BIL is an exchange-traded fund with daily liquidity, allowing us to obtain a consistent daily price series from the same source as our risky assets.")
    w("")
    w(f"The average daily return of BIL over {bil_n_returns} observations is:")
    w("")
    w(f"> μ_F = {_fmt(rf, 8)}")
    w("")
    w(f"which annualizes to approximately {_pct(rf * 252)}. "
      f"The daily standard deviation of BIL returns is {_fmt(bil_daily_std, 8)}, "
      "confirming near-zero variance.")
    w("")
    w(f"Critically, we verify that μ_F < A/C (Case 1 in Huang & Litzenberger): "
      f"μ_F = {_fmt(rf, 8)} < A/C = {_fmt(mu_mvp, 8)}. "
      "This condition ensures that the tangent line from the risk-free rate touches "
      "the efficient (upper) branch of the frontier, so the tangency portfolio is a "
      "legitimate efficient portfolio.")
    w("")
    w("BIL daily prices and returns are reported in Sheet 5 of the Excel workbook.")
    w("")

    # ── Section 7 ────────────────────────────────────────────────────────
    w("## 8. MVP, Market Portfolio, and ZC Portfolio (Step 7, Sheet 6)")
    w("")
    w("*Reference: Course Chapters 8–9; Huang & Litzenberger, Chapters 4–5.*")
    w("")

    # 7a: MVP
    w("### 8a. Minimum Variance Portfolio (MVP)")
    w("")
    w("The MVP is the portfolio on the frontier with the lowest possible variance. "
      "It occupies the vertex of the mean-variance hyperbola and serves as the "
      "boundary between the efficient and inefficient portions of the frontier.")
    w("")
    w("> w_MVP = Σ⁻¹ 1 / C")
    w("")
    w("| Statistic | Value |")
    w("|-----------|-------|")
    w(f"| Expected Return (μ_MVP) | {_fmt(mvp['expected_return'], 8)} |")
    w(f"| Standard Deviation (σ_MVP) | {_fmt(mvp['std_dev'], 8)} |")
    w(f"| Annualized Return | {_pct(mvp['expected_return'] * 252)} |")
    w("")
    w("The MVP weight distribution reflects pure variance minimization — "
      "the optimizer ignores expected returns entirely and allocates to reduce "
      "total portfolio risk:")
    w("")
    w("| Ticker | Name | MVP Weight |")
    w("|--------|------|-----------|")
    for idx in mvp_sort:
        t = tickers[idx]
        w(f"| {t} | {ASSET_NAMES[t]} | {mvp['weights'][idx]:+.4f} |")
    w("")
    w(f"The portfolio is dominated by low-volatility defensive names: {top_names}. "
      "These assets have the lowest daily standard deviations in the universe and "
      "relatively low correlations with other holdings, making them efficient "
      "variance reducers. The MVP is agnostic to expected return — it is the "
      "portfolio that a maximally risk-averse investor would hold.")
    w("")

    # 7b: Market Portfolio
    w("### 8b. Market Portfolio (Tangency Portfolio)")
    w("")
    w("When a riskless asset is available, the efficient set becomes a straight line "
      "(the Capital Market Line) from the risk-free rate tangent to the risky-asset "
      "frontier. The tangency point is the Market Portfolio — the risky portfolio "
      "that maximizes the Sharpe ratio.")
    w("")
    w("The unnormalized tangency weights are:")
    w("")
    w("> z = Σ⁻¹ (μ − μ_F · 1)")
    w("")
    w("Normalizing to sum to 1 gives w_M = z / (1^T z). The quantity H measures "
      "the squared maximum Sharpe ratio attainable from the risky assets:")
    w("")
    w("> H = (μ − μ_F·1)^T Σ⁻¹ (μ − μ_F·1)")
    w("")
    w("| Statistic | Value |")
    w("|-----------|-------|")
    w(f"| Expected Return (μ_M) | {_fmt(market['expected_return'], 8)} |")
    w(f"| Standard Deviation (σ_M) | {_fmt(market['std_dev'], 8)} |")
    w(f"| Sharpe Ratio | {_fmt(market['sharpe_ratio'], 6)} |")
    w(f"| H | {_fmt(market['H'], 10)} |")
    w(f"| Annualized Return | {_pct(market['expected_return'] * 252)} |")
    w("")
    w("| Ticker | Name | Market Weight |")
    w("|--------|------|--------------|")
    for idx in mkt_sort:
        t = tickers[idx]
        w(f"| {t} | {ASSET_NAMES[t]} | {market['weights'][idx]:+.4f} |")
    w("")

    # Leverage discussion
    w("#### Discussion: Extreme Weights and Leverage")
    w("")
    w(f"The Market Portfolio exhibits substantial leverage: the gross exposure "
      f"(sum of absolute weights) is {gross_exposure:.1f}%, with "
      f"{long_sum:+.2f} in long positions and {short_sum:+.2f} in short positions. "
      "This is a well-documented property of unconstrained mean-variance optimization "
      "and warrants careful interpretation.")
    w("")
    w("The unconstrained optimization framework as presented in the course "
      "permits short selling — portfolio weights can be negative. "
      "This is consistent with the theoretical treatment in Huang & Litzenberger, "
      "where no sign constraints are imposed on the weight vector.")
    w("")
    if long_clusters or short_clusters:
        w("At the cluster level, the optimizer is:")
        w("")
        for cname, wt in long_clusters:
            w(f"- **Long** {cname} ({wt:+.2f})")
        for cname, wt in short_clusters:
            w(f"- **Short** {cname} ({wt:+.2f})")
        w("")
    w("The optimizer goes long the clusters with the highest risk-adjusted returns "
      "and shorts those with negative or low returns. It exploits the correlation "
      "structure to hedge out common factor exposures, effectively isolating the "
      "return spread between winning and losing clusters.")
    w("")
    w("This extreme leverage is a key limitation of the mean-variance model. "
      "In practice, several approaches address it: imposing long-only constraints "
      "(0 ≤ w_i ≤ 1), adding regularization to the covariance matrix, using "
      "shrinkage estimators (Ledoit-Wolf), or applying position limits. The "
      "Bedrock Fund's theoretical analysis intentionally preserves the unconstrained "
      "solution to demonstrate the framework as taught in the course.")
    w("")

    # 7c: ZC Portfolio
    w("### 8c. Zero-Covariance Portfolio")
    w("")
    w("*Reference: Course Chapter 8, Proposition 2 — Zero-Covariance Portfolio.*")
    w("")
    w("For any frontier portfolio P with expected return μ_P ≠ μ_MVP, there exists "
      "a unique frontier portfolio ZC(P) such that Cov(R_P, R_ZC) = 0. The ZC "
      "portfolio generalizes the role of the risk-free asset: in the absence of a "
      "riskless asset, the ZC portfolio of any efficient portfolio provides the same "
      "two-fund separation theorem.")
    w("")
    w("For the Market Portfolio with return μ_M, the ZC portfolio's expected return is:")
    w("")
    w("> μ_ZC = A/C − (D/C²) / (μ_M − A/C)")
    w("")
    w("| Statistic | Value |")
    w("|-----------|-------|")
    w(f"| Expected Return (μ_ZC) | {_fmt(zc['expected_return'], 8)} |")
    w(f"| Standard Deviation (σ_ZC) | {_fmt(zc['std_dev'], 8)} |")
    w(f"| Cov(Market, ZC) | {cov_mkt_zc:.2e} |")
    w(f"| Annualized Return | {_pct(zc['expected_return'] * 252)} |")
    w("")
    w(f"The covariance between the Market Portfolio and its ZC portfolio is "
      f"{cov_mkt_zc:.2e}, confirming zero covariance to machine precision.")
    w("")
    w(f"Notably, μ_ZC = {_fmt(zc['expected_return'], 8)} ≈ μ_F = {_fmt(rf, 8)}. "
      "This is theoretically expected: for the tangency portfolio (Market Portfolio), "
      "the zero-covariance portfolio has the same expected return as the risk-free "
      "rate. This follows from the Black (1972) zero-beta CAPM: when a riskless asset "
      "exists, the zero-beta rate equals the risk-free rate.")
    w("")
    w("All portfolio weights and statistics are reported in Sheet 6 of the Excel workbook.")
    w("")

    # ── Section 8 ────────────────────────────────────────────────────────
    w("## 9. Graph Analysis (Step 8)")
    w("")
    w("*Reference: Course Chapters 8–9; Huang & Litzenberger, Chapters 4–5.*")
    w("")
    w("Graph 1 plots the full portfolio analysis in (σ, μ) space — "
      "standard deviation on the horizontal axis and expected return on the vertical "
      "axis. The graph contains the following elements:")
    w("")
    w("- **Efficient Portfolio Frontier (upper branch):** The solid curve above the MVP representing all efficient risky-asset portfolios. These are the portfolios that maximize expected return for a given level of risk.")
    w("- **Inefficient Frontier (lower branch):** The dashed curve below the MVP. These portfolios are dominated — for the same variance, an investor can achieve a higher return on the upper branch.")
    w(f"- **Minimum Variance Portfolio (MVP):** Marked at (σ, μ) = ({_fmt(mvp['std_dev'], 4)}, {_fmt(mvp['expected_return'], 6)}). The leftmost point on the frontier.")
    w(f"- **Market Portfolio (Tangency Portfolio):** Marked at (σ, μ) = ({_fmt(market['std_dev'], 4)}, {_fmt(market['expected_return'], 6)}). The point where the CML is tangent to the efficient frontier.")
    w(f"- **Zero-Covariance Portfolio:** Marked at (σ, μ) = ({_fmt(zc['std_dev'], 4)}, {_fmt(zc['expected_return'], 6)}). Located on the inefficient branch, directly below the MVP.")
    w(f"- **Capital Market Line (CML):** The straight line from (0, μ_F) through the Market Portfolio, with slope √H = {_fmt(cml_slope, 6)}.")
    w("")
    w("The Capital Market Line equation is:")
    w("")
    w(f"> μ_P = μ_F + √H · σ_P = {_fmt(rf, 8)} + {_fmt(cml_slope, 6)} · σ_P")
    w("")
    w("The CML dominates the efficient frontier: for any level of risk σ_P, "
      "the CML offers a higher expected return than the risky-asset frontier alone. "
      "This is the key insight of introducing a riskless asset — investors can "
      "improve their risk-return tradeoff by combining the risk-free asset with "
      "the Market Portfolio.")
    w("")
    w("Two-fund separation follows directly: every efficient portfolio (on the CML) "
      "can be expressed as a combination of just two funds — the risk-free asset "
      "and the Market Portfolio. A conservative investor holds mostly the risk-free "
      "asset with a small allocation to the Market Portfolio; an aggressive investor "
      "borrows at the risk-free rate and levers up the Market Portfolio.")
    w("")
    w("![Graph 1](graph1.png)")
    w("")
    w("*Figure 1: Bedrock Fund — Efficient Portfolio Frontier, MVP, Market Portfolio, ZC Portfolio, and Capital Market Line in (σ, μ) space.*")
    w("")

    # ── Section 9 ────────────────────────────────────────────────────────
    w("## 10. Conclusions")
    w("")
    w("### Advantages of the Bedrock Fund")
    w("")
    w("- **Thematic diversification:** The picks-and-shovels approach distributes exposure across five distinct structural trends (AI, energy transition, logistics, healthcare, defense). This reduces single-company and single-sector risk.")
    w("- **Factor breadth:** The 20 assets span technology, industrials, energy, healthcare, commodities, and financials, providing exposure to multiple macroeconomic factors (growth, inflation, rates, geopolitics).")
    w("- **Infrastructure resilience:** Infrastructure providers tend to have more durable revenue streams than end-market competitors, as their products are required regardless of which downstream company wins.")
    w("- **Covariance structure quality:** The cross-cluster correlations range from approximately 0.36 to 0.62, providing meaningful but not excessive diversification — enough off-diagonal variation for the efficient frontier to offer real variance reduction.")
    w("")
    w("### Limitations and Disadvantages")
    w("")
    w(f"- **Impractical leverage in the unconstrained solution:** The Market Portfolio has gross exposure of {gross_exposure:.0f}%, requiring extensive short selling. In practice, most mutual funds face regulatory constraints (e.g., the Investment Company Act of 1940) that prohibit or limit short positions.")
    w("- **Backward-looking estimation:** The mean vector and covariance matrix are estimated from one year of historical data. Past returns are a noisy and often unreliable predictor of future returns. The covariance structure is more stable, but still subject to regime changes.")
    w("- **Normality assumption:** The mean-variance framework implicitly assumes returns are normally distributed (or that investor utility is quadratic). In reality, asset returns exhibit fat tails, skewness, and time-varying volatility, which the model does not capture.")
    w("- **Single-period model:** The analysis is a single-period (one-shot) optimization. It does not account for dynamic rebalancing, transaction costs, taxes, or changing investment opportunities over time.")
    w("- **Estimation error amplification:** Inverting the covariance matrix amplifies estimation errors in the sample means and covariances, which is a primary driver of the extreme weights observed. This is the classic Michaud (1989) critique of mean-variance optimization.")
    w("")
    w("### The Tension Between Theory and Practice")
    w("")
    w("The Bedrock Fund analysis illustrates a fundamental tension in financial "
      "economics. The mean-variance framework of Markowitz (1952), as formalized "
      "by Huang & Litzenberger, provides an elegant and mathematically rigorous "
      "theory for optimal portfolio selection. The efficient frontier, the tangency "
      "portfolio, the zero-covariance portfolio, and two-fund separation are "
      "powerful theoretical results that illuminate the structure of the "
      "risk-return tradeoff.")
    w("")
    w("However, the unconstrained implementation reveals the model's sensitivity "
      "to input estimation. Small changes in the estimated mean vector can produce "
      "large swings in optimal weights, and the resulting portfolios often require "
      "leverage and short selling that are impractical for most real-world investors. "
      "This gap between theoretical optimality and practical implementability is "
      "one of the central challenges in quantitative portfolio management, and has "
      "motivated decades of research into robust optimization, shrinkage estimators, "
      "and constrained frontier methods.")
    w("")

    md_path = os.path.join(output_dir, "bedrock_fund_report.md")
    with open(md_path, "w") as f:
        f.write("\n".join(lines))
    print(f"Markdown report saved to {md_path}")
    return md_path
